"""
Resume text extraction utility.

Implements: PRD §8.1 (Upload/parse resume (PDF/DOCX) → structured profile).
Roadmap: Epic 2 - Resume Ingestion & Profile Extraction, Story 2 - Structured
Profile Extraction, Task 1.

Extracts raw text from resume file bytes (PDF or DOCX). This module owns only
text extraction — no LLM calls, no persistence, no HTTP concerns (Single
Responsibility, per docs/coding_guidelines.md). Structured-field extraction
from this text is handled separately in
`app/agents/resume_parser/parser_agent.py` (Epic 2, Story 2, Task 2).
"""

from __future__ import annotations

import io
from abc import ABC, abstractmethod

import docx
from pypdf import PdfReader


class ResumeExtractionError(Exception):
    """Raised when resume text cannot be extracted from the given content."""


class ResumeTextExtractor(ABC):
    """Abstraction over extracting plain text from a resume file format."""

    @abstractmethod
    def supports(self, extension: str) -> bool:
        """Return True if this extractor can handle the given file extension."""
        raise NotImplementedError

    @abstractmethod
    def extract_text(self, content: bytes) -> str:
        """Extract plain text from raw file bytes.

        Raises:
            ResumeExtractionError: if text cannot be extracted.
        """
        raise NotImplementedError


class PdfResumeExtractor(ResumeTextExtractor):
    """Extracts text from PDF resumes using pypdf."""

    def supports(self, extension: str) -> bool:
        return extension.lower() == ".pdf"

    def extract_text(self, content: bytes) -> str:
        try:
            reader = PdfReader(io.BytesIO(content))
        except Exception as exc:  # noqa: BLE001 - normalize third-party failure
            raise ResumeExtractionError("Failed to open PDF content for reading.") from exc

        pages_text: list[str] = []
        for page_number, page in enumerate(reader.pages):
            try:
                text = page.extract_text() or ""
            except Exception as exc:  # noqa: BLE001 - normalize per-page failure
                raise ResumeExtractionError(
                    f"Failed to extract text from PDF page {page_number}."
                ) from exc
            if text:
                pages_text.append(text)

        full_text = "\n".join(pages_text).strip()
        if not full_text:
            raise ResumeExtractionError(
                "PDF contained no extractable text (it may be a scanned image)."
            )
        return full_text


class DocxResumeExtractor(ResumeTextExtractor):
    """Extracts text from DOCX resumes using python-docx."""

    def supports(self, extension: str) -> bool:
        return extension.lower() == ".docx"

    def extract_text(self, content: bytes) -> str:
        try:
            document = docx.Document(io.BytesIO(content))
        except Exception as exc:  # noqa: BLE001 - normalize third-party failure
            raise ResumeExtractionError("Failed to open DOCX content for reading.") from exc

        paragraphs = [p.text for p in document.paragraphs if p.text and p.text.strip()]

        table_cells: list[str] = []
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.text.strip():
                        table_cells.append(cell.text.strip())

        full_text = "\n".join([*paragraphs, *table_cells]).strip()
        if not full_text:
            raise ResumeExtractionError("DOCX contained no extractable text.")
        return full_text


class ResumeTextExtractorRegistry:
    """Dispatches to the appropriate extractor based on file extension.

    Adding support for a new resume file format means adding a new
    `ResumeTextExtractor` implementation and registering it here — no
    existing extractor is modified (Open/Closed, per docs/architecture.md).
    """

    def __init__(self, extractors: list[ResumeTextExtractor] | None = None) -> None:
        self._extractors: list[ResumeTextExtractor] = extractors or [
            PdfResumeExtractor(),
            DocxResumeExtractor(),
        ]

    def extract_text(self, content: bytes, extension: str) -> str:
        for extractor in self._extractors:
            if extractor.supports(extension):
                return extractor.extract_text(content)

        raise ResumeExtractionError(
            f"No extractor registered for file extension '{extension}'."
        )